"""文件解析服务，支持 PDF / Word / Excel / Markdown / TXT / HTML"""
import os
from pathlib import Path
from typing import Optional

from app.config import UPLOAD_PATH, FILE_PARSER_MAP, ALLOWED_FILE_TYPES, MAX_FILE_SIZE
from app.utils.response import APIError


def get_file_type(filename: str) -> str:
    """从文件名获取类型（扩展名，不带点），未知返回空字符串"""
    ext = Path(filename).suffix.lower()
    return FILE_PARSER_MAP.get(ext, "")


def parse_file(file_path: str, file_type: str) -> str:
    """解析文件，返回提取的纯文本内容

    Args:
        file_path: 文件绝对路径
        file_type: 文件类型(pdf/docx/xlsx/markdown)
    """
    if not os.path.exists(file_path):
        raise APIError(f"文件不存在: {file_path}", status_code=404)

    try:
        if file_type == "pdf":
            return _parse_pdf(file_path)
        if file_type == "docx":
            return _parse_docx(file_path)
        if file_type == "xlsx":
            return _parse_xlsx(file_path)
        if file_type in ("markdown", "md"):
            return _parse_markdown(file_path)
        if file_type == "text":
            return _parse_text(file_path)
        if file_type == "html":
            return _parse_html(file_path)
        raise APIError(f"不支持的文件类型: {file_type}")
    except APIError:
        raise
    except Exception as e:
        raise APIError(f"文件解析失败: {str(e)}")


def _parse_pdf(file_path: str) -> str:
    """解析PDF文件"""
    from PyPDF2 import PdfReader
    reader = PdfReader(file_path)
    texts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        texts.append(text)
    return "\n".join(texts).strip()


def _parse_docx(file_path: str) -> str:
    """解析Word docx文件"""
    from docx import Document
    doc = Document(file_path)
    texts = []
    for para in doc.paragraphs:
        if para.text:
            texts.append(para.text)
    # 表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                texts.append(" | ".join(cells))
    return "\n".join(texts).strip()


def _parse_xlsx(file_path: str) -> str:
    """解析Excel xlsx文件"""
    from openpyxl import load_workbook
    wb = load_workbook(file_path, data_only=True, read_only=True)
    texts = []
    for sheet in wb.worksheets:
        texts.append(f"# {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                texts.append(" | ".join(cells))
    wb.close()
    return "\n".join(texts).strip()


def _parse_markdown(file_path: str) -> str:
    """读取Markdown文件"""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read().strip()


def _parse_text(file_path: str) -> str:
    """读取纯文本文件，自动尝试多种编码"""
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030", "latin-1"):
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read().strip()
        except (UnicodeDecodeError, UnicodeError):
            continue
    # 兜底：utf-8 忽略无法解码的字符
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()


def _parse_html(file_path: str) -> str:
    """解析HTML文件，提取纯文本内容"""
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030", "latin-1"):
        try:
            with open(file_path, "r", encoding=enc) as f:
                html_content = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    else:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            html_content = f.read()

    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")
        # 移除 script 和 style 标签
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        lines = [line.strip() for line in text.splitlines()]
        return "\n".join(line for line in lines if line).strip()
    except ImportError:
        # bs4 未安装时用正则简单去标签
        import re
        text = re.sub(r"<script[^>]*>.*?</script>", "", html_content, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<[^>]+>", "\n", text)
        text = re.sub(r"&nbsp;", " ", text)
        text = re.sub(r"&amp;", "&", text)
        text = re.sub(r"&lt;", "<", text)
        text = re.sub(r"&gt;", ">", text)
        lines = [line.strip() for line in text.splitlines()]
        return "\n".join(line for line in lines if line).strip()


def _sanitize_filename(filename: str) -> str:
    """清理文件名：去除路径分隔符、危险字符，仅保留安全字符

    防止路径穿越攻击（如 ../../etc/passwd）和非法文件名。
    """
    import re
    # 仅取 basename（去除任何路径前缀）
    name = os.path.basename(filename.replace("\\", "/"))
    # 去除 Windows 保留字符 <>:"/\\|?* 和控制字符
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', name)
    # 去除以点开头（隐藏文件）和以点结尾（Windows 兼容）
    name = name.lstrip('.').rstrip('. ')
    # 限制长度（保留扩展名）
    if '.' in name:
        stem, ext = name.rsplit('.', 1)
        ext = '.' + ext[:10]
    else:
        stem, ext = name, ''
    max_stem = 200 - len(ext)
    name = stem[:max_stem] + ext
    # 兜底：空文件名
    if not name or name == ext:
        name = f"upload{ext}"
    return name


def save_upload_file(upload_file, filename: str) -> str:
    """保存上传文件到 UPLOAD_PATH，返回存储路径

    Args:
        upload_file: FastAPI UploadFile 对象
        filename: 目标文件名
    """
    file_type = get_file_type(filename)
    if not file_type:
        raise APIError(f"不支持的文件类型，允许: {', '.join(FILE_PARSER_MAP.keys())}")

    # 清理文件名（防路径穿越 + 防危险字符）+ 加时间戳前缀防冲突
    import time
    clean_name = _sanitize_filename(filename)
    safe_name = f"{int(time.time())}_{clean_name}"
    target_dir = Path(UPLOAD_PATH)
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / safe_name

    # 流式分片写入：避免大文件全量读入内存导致 OOM
    written = 0
    chunk_size = 1024 * 1024  # 1MB
    try:
        with open(target_path, "wb") as f:
            while True:
                chunk = upload_file.file.read(chunk_size)
                if not chunk:
                    break
                written += len(chunk)
                if written > MAX_FILE_SIZE:
                    f.close()
                    # 删除已写入的部分文件
                    try:
                        os.remove(target_path)
                    except OSError:
                        pass
                    raise APIError(f"文件大小超过限制({MAX_FILE_SIZE // 1024 // 1024}MB)")
                f.write(chunk)
    finally:
        upload_file.file.close()

    return str(target_path)
