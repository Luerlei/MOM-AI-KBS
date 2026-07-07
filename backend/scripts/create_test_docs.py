#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
create_test_docs.py
===================

生成 6 个预制测试文档，用于测试知识库的文档解析与检索功能。

输出目录：backend/data/test_docs/

生成文件：
    1. 设备故障代码手册.pdf      (PDF,  reportlab / fpdf2 / txt 回退)
    2. 设备保养计划规程.docx     (Word, python-docx)
    3. 数控加工工艺规程.docx     (Word, python-docx)
    4. 产品质量检验标准.xlsx     (Excel, openpyxl)
    5. 仓库管理制度.pdf          (PDF,  reportlab / fpdf2 / txt 回退)
    6. 设备操作经验记录.md        (Markdown, 直接写入)

运行方式：
    cd backend
    python scripts/create_test_docs.py

PDF 生成策略（按优先级）：
    1. reportlab        （优先，支持嵌入中文字体）
    2. fpdf2            （次选，支持嵌入中文字体）
    3. 纯标准库生成器   （无第三方依赖，嵌入 Windows 系统 TTF 字体，
                         始终可用，附带 ToUnicode CMap 以便 PyPDF2 提取文本）
    4. 自动尝试 pip install fpdf2
    5. 以上均不可用 → 回退为 .txt 文件（保证脚本始终能运行成功）
"""

import os
import sys
import struct
import subprocess

# --------------------------------------------------------------------------- #
# 路径计算：基于脚本自身位置，与调用时的 cwd 无关
# --------------------------------------------------------------------------- #
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BACKEND_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BACKEND_DIR, "data", "test_docs")


# --------------------------------------------------------------------------- #
# 工具函数
# --------------------------------------------------------------------------- #
def ensure_dir(path):
    """确保目录存在。"""
    os.makedirs(path, exist_ok=True)


def find_cjk_font():
    """
    在 Windows 系统字体目录中查找一个可用的中文 TTF/TTC 字体文件路径。
    优先返回单文件 TTF（兼容性最好），其次 TTC。
    找不到返回 None。
    """
    candidates_ttf = [
        "simhei.ttf",   # 黑体（单文件，兼容性最佳）
        "simsun.ttf",
        "msyh.ttf",     # 微软雅黑
        "Deng.ttf",     # 等线
        "simkai.ttf",
    ]
    candidates_ttc = [
        "msyh.ttc",     # 微软雅黑（集合）
        "simsun.ttc",   # 宋体（集合）
        "msyhbd.ttc",
    ]
    fonts_dir = os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts")
    for name in candidates_ttf + candidates_ttc:
        p = os.path.join(fonts_dir, name)
        if os.path.exists(p):
            return p
    return None


# --------------------------------------------------------------------------- #
# PDF 生成
# --------------------------------------------------------------------------- #
def _pdf_with_reportlab(path, lines, font_path):
    """使用 reportlab 生成 PDF（嵌入中文字体）。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 注册中文字体
    font_name = "Helvetica"  # 默认回退（无法显示中文）
    if font_path:
        try:
            if font_path.lower().endswith(".ttc"):
                pdfmetrics.registerFont(TTFont("CJK", font_path, subfontIndex=0))
            else:
                pdfmetrics.registerFont(TTFont("CJK", font_path))
            font_name = "CJK"
        except Exception as e:
            print(f"  [warn] reportlab 注册中文字体失败: {e}，将使用默认字体（中文可能无法显示）")

    width, height = A4
    c = canvas.Canvas(path, pagesize=A4)
    c.setFont(font_name, 11)
    left, top, line_height = 40, height - 40, 16
    y = top
    for raw in lines:
        text = raw.rstrip("\n")
        if y < 40:  # 分页
            c.showPage()
            c.setFont(font_name, 11)
            y = top
        c.drawString(left, y, text)
        y -= line_height
    c.save()


def _pdf_with_fpdf2(path, lines, font_path):
    """使用 fpdf2 生成 PDF（嵌入中文字体）。"""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    family = "Helvetica"
    if font_path:
        try:
            pdf.add_font("CJK", "", font_path, uni=True)
            family = "CJK"
        except Exception as e:
            print(f"  [warn] fpdf2 加载中文字体失败: {e}，将使用默认字体（中文可能无法显示）")
    pdf.set_font(family, size=11)

    for raw in lines:
        text = raw.rstrip("\n")
        try:
            pdf.multi_cell(0, 7, text)
        except TypeError:
            # 兼容新版 fpdf2 的 multi_cell 签名
            try:
                pdf.multi_cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")
            except Exception:
                pdf.cell(0, 7, text, ln=1)
        except Exception as e:
            # 极端情况下用 latin-1 替换无法编码的字符
            safe = text.encode("latin-1", "replace").decode("latin-1")
            pdf.cell(0, 7, safe, ln=1)

    # fpdf2 的 output() 无参数时返回 bytearray
    data = pdf.output()
    with open(path, "wb") as f:
        f.write(bytes(data))


# --------------------------------------------------------------------------- #
# 纯标准库 PDF 生成（无第三方依赖）
# 通过解析 TTF 的 cmap 表将 Unicode 映射到字形 ID，嵌入完整 TTF 字体，
# 并写入 ToUnicode CMap，使 PyPDF2 等工具能正确提取中文文本。
# --------------------------------------------------------------------------- #
def _ttf_table_dir(data):
    """解析 TTF/OTF 表目录，返回 {tag: (offset, length)}。"""
    if len(data) < 12:
        raise ValueError("not a valid font file")
    sf_version, num_tables = struct.unpack(">IH", data[0:6])
    tables = {}
    pos = 12
    for _ in range(num_tables):
        tag = data[pos:pos + 4].decode("latin-1")
        _checksum, offset, length = struct.unpack(">III", data[pos + 4:pos + 16])
        tables[tag] = (offset, length)
        pos += 16
    return tables


def _cmap_lookup(data, cmap_offset):
    """解析 cmap 表，返回一个 char -> glyphID 的查找函数。"""
    _version, num_sub = struct.unpack(">HH", data[cmap_offset:cmap_offset + 4])
    best_off, best_pri = None, -1
    pos = cmap_offset + 4
    for _ in range(num_sub):
        pid, psid, sub_off = struct.unpack(">HHI", data[pos:pos + 8])
        pos += 8
        priority = {(3, 10): 3, (0, 4): 2, (3, 1): 1, (0, 3): 0}.get((pid, psid), -2)
        if priority > best_pri:
            best_pri, best_off = priority, cmap_offset + sub_off
    if best_off is None:
        return lambda c: 0
    fmt = struct.unpack(">H", data[best_off:best_off + 2])[0]
    if fmt == 4:
        return _cmap_fmt4(data, best_off)
    if fmt == 12:
        return _cmap_fmt12(data, best_off)
    return lambda c: 0


def _cmap_fmt4(data, offset):
    _fmt, _length, _language, seg_x2 = struct.unpack(">HHHH", data[offset:offset + 8])
    seg_count = seg_x2 // 2
    p = offset + 14  # 8 + 6 (searchRange/entrySelector/rangeShift)
    end_codes = struct.unpack(">" + "H" * seg_count, data[p:p + seg_count * 2]); p += seg_count * 2
    p += 2  # reservedPad
    start_codes = struct.unpack(">" + "H" * seg_count, data[p:p + seg_count * 2]); p += seg_count * 2
    id_deltas = struct.unpack(">" + "h" * seg_count, data[p:p + seg_count * 2]); p += seg_count * 2
    id_range_offsets = struct.unpack(">" + "H" * seg_count, data[p:p + seg_count * 2]); p += seg_count * 2
    gid_array_start = p
    idro_start = gid_array_start - seg_count * 2

    def lookup(c):
        for i in range(seg_count):
            if end_codes[i] >= c:
                if start_codes[i] > c:
                    return 0
                if id_range_offsets[i] == 0:
                    return (c + id_deltas[i]) & 0xFFFF
                addr = idro_start + i * 2 + id_range_offsets[i] + 2 * (c - start_codes[i])
                if addr + 2 > len(data):
                    return 0
                g_raw = struct.unpack(">H", data[addr:addr + 2])[0]
                if g_raw == 0:
                    return 0
                return (g_raw + id_deltas[i]) & 0xFFFF
        return 0
    return lookup


def _cmap_fmt12(data, offset):
    _fmt, _reserved, _length, _language, num_groups = struct.unpack(">HHIII", data[offset:offset + 16])
    groups = []
    p = offset + 16
    for _ in range(num_groups):
        sc, ec, sg = struct.unpack(">III", data[p:p + 12])
        groups.append((sc, ec, sg))
        p += 12
    import bisect
    starts = [g[0] for g in groups]

    def lookup(c):
        i = bisect.bisect_right(starts, c) - 1
        if i < 0:
            return 0
        sc, ec, sg = groups[i]
        return sg + (c - sc) if sc <= c <= ec else 0
    return lookup


def _pdf_with_stdlib(path, lines, font_path):
    """纯标准库生成 PDF：嵌入 Windows 系统 TTF 字体，中文可正常显示与提取。"""
    with open(font_path, "rb") as f:
        fdata = f.read()
    tables = _ttf_table_dir(fdata)
    if "cmap" not in tables or "head" not in tables or "hhea" not in tables:
        raise ValueError("字体缺少必要的表 (cmap/head/hhea)")

    glyph_for = _cmap_lookup(fdata, tables["cmap"][0])

    head = tables["head"][0]
    upem = struct.unpack(">H", fdata[head + 18:head + 20])[0]
    x_min, y_min, x_max, y_max = struct.unpack(">hhhh", fdata[head + 36:head + 44])

    hhea = tables["hhea"][0]
    ascent = struct.unpack(">h", fdata[hhea + 4:hhea + 6])[0]
    descent = struct.unpack(">h", fdata[hhea + 6:hhea + 8])[0]
    num_hmetrics = struct.unpack(">H", fdata[hhea + 34:hhea + 36])[0]

    hmtx = tables["hmtx"][0] if "hmtx" in tables else 0

    def width_1000(gid):
        if gid < 0:
            return 1000
        if num_hmetrics and hmtx:
            if gid < num_hmetrics:
                w = struct.unpack(">H", fdata[hmtx + gid * 4 + 2:hmtx + gid * 4 + 4])[0]
            else:
                w = struct.unpack(">H", fdata[hmtx + (num_hmetrics - 1) * 4 + 2:hmtx + (num_hmetrics - 1) * 4 + 4])[0]
            return round(w * 1000 / upem) if upem else 1000
        return 1000

    # 收集所有使用到的字形与宽度，构建 W 数组
    used_gids = {}
    gid_to_char = {}
    for line in lines:
        for ch in line:
            c = ord(ch)
            g = glyph_for(c)
            used_gids.setdefault(g, width_1000(g))
            if g and g not in gid_to_char:
                gid_to_char[g] = c
    w_array = " ".join(f"{g} [{w}]" for g, w in sorted(used_gids.items()))

    # ToUnicode CMap（让 PyPDF2 能把字形 ID 反解为 Unicode 文本）
    tu_lines = [
        "/CIDInit /ProcSet findresource begin",
        "12 dict begin",
        "begincmap",
        "/CIDSystemInfo <</Registry (Adobe)/Ordering (UCS)/Supplement 0>> def",
        "/CMapName /Adobe-Identity-UCS def",
        "/CMapType 2 def",
        "1 begincodespacerange",
        "<0000> <FFFF>",
        "endcodespacerange",
    ]
    items = sorted(gid_to_char.items())
    i = 0
    while i < len(items):
        chunk = items[i:i + 100]
        tu_lines.append(f"{len(chunk)} beginbfchar")
        for g, c in chunk:
            tu_lines.append(f"<{g:04X}> <{c:04X}>")
        tu_lines.append("endbfchar")
        i += 100
    tu_lines += ["endcmap", "CMapName currentdict /CMap defineresource pop", "end", "end"]
    tu_bytes = "\n".join(tu_lines).encode("ascii")

    # 页面参数与分页
    page_w, page_h = 595, 842
    left, top, line_h, font_size = 40, page_h - 40, 16, 11
    lines_per_page = int((page_h - 80) // line_h)
    page_chunks = [lines[i:i + lines_per_page] for i in range(0, len(lines), lines_per_page)] or [[]]
    num_pages = len(page_chunks)

    # 预分配对象编号：1 Catalog, 2 Pages, 3 Type0, 4 CIDFont, 5 FontDesc,
    #                6 FontFile2, 7 ToUnicode, 8.. pages/contents
    def page_obj(k):
        return 8 + 2 * k

    def contents_obj(k):
        return 9 + 2 * k

    kids = " ".join(f"{page_obj(k)} 0 R" for k in range(num_pages))

    catalog = "<</Type/Catalog/Pages 2 0 R>>"
    pages_obj = f"<</Type/Pages/Kids[{kids}]/Count {num_pages}>>"
    type0 = ("<</Type/Font/Subtype/Type0/BaseFont/SimHei/Encoding/Identity-H/"
             "DescendantFonts[4 0 R]/ToUnicode 7 0 R>>")
    cidfont = ("<</Type/Font/Subtype/CIDFontType2/BaseFont/SimHei/CIDToGIDMap/Identity/"
               f"FontDescriptor 5 0 R/DW 1000/W[{w_array}]>>")
    fontdesc = ("<</Type/FontDescriptor/FontName/SimHei/Flags 4/"
                f"FontBBox[{round(x_min*1000/upem)} {round(y_min*1000/upem)} "
                f"{round(x_max*1000/upem)} {round(y_max*1000/upem)}]/ItalicAngle 0/"
                f"Ascent {round(ascent*1000/upem)}/Descent {round(descent*1000/upem)}/"
                f"CapHeight {round(ascent*1000/upem)}/StemV 80/FontFile2 6 0 R>>")
    fontfile = (b"<</Length " + str(len(fdata)).encode() + b"/Length1 "
                + str(len(fdata)).encode() + b">>stream\n" + fdata + b"\nendstream")
    tounicode = (b"<</Length " + str(len(tu_bytes)).encode() + b">>stream\n"
                 + tu_bytes + b"\nendstream")

    page_bodies, content_bodies = [], []
    for k, page_lines in enumerate(page_chunks):
        parts = []
        y = top
        for line in page_lines:
            text = line.rstrip("\n")
            gids = [glyph_for(ord(ch)) for ch in text]
            hexs = "".join(f"{g:04X}" for g in gids)
            if hexs:
                parts.append(f"BT /F1 {font_size} Tf {left} {y} Td <{hexs}> Tj ET")
            else:
                parts.append(f"BT /F1 {font_size} Tf {left} {y} Td ET")
            y -= line_h
        content_bytes = "\n".join(parts).encode("ascii")
        # 内容流对象必须包装为流：<< /Length N >> stream ... endstream
        content_bodies.append(
            b"<</Length " + str(len(content_bytes)).encode() + b">>stream\n"
            + content_bytes + b"\nendstream"
        )
        page_bodies.append(
            f"<</Type/Page/Parent 2 0 R/MediaBox[0 0 {page_w} {page_h}]/"
            f"Resources<</Font<</F1 3 0 R>>>>/Contents {contents_obj(k)} 0 R>>"
        )

    all_objs = [catalog, pages_obj, type0, cidfont, fontdesc, fontfile, tounicode]
    for k in range(num_pages):
        all_objs.append(page_bodies[k])
        all_objs.append(content_bodies[k])

    out = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    offsets = []
    for idx, body in enumerate(all_objs, 1):
        offsets.append(len(out))
        b = body if isinstance(body, bytes) else body.encode("latin-1")
        out += f"{idx} 0 obj\n".encode() + b + b"\nendobj\n"
    xref_pos = len(out)
    n = len(all_objs)
    out += f"xref\n0 {n + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += (f"trailer\n<</Size {n + 1}/Root 1 0 R>>\nstartxref\n{xref_pos}\n%%EOF").encode()
    with open(path, "wb") as f:
        f.write(out)


def write_pdf(filename, lines):
    """
    生成 PDF 文件。
    按优先级尝试 reportlab → fpdf2 → 纯标准库 → 自动安装 fpdf2 → 回退 txt。
    返回 (实际文件路径, 使用的引擎名称)。
    """
    path = os.path.join(OUTPUT_DIR, filename)
    font_path = find_cjk_font()

    # 1) reportlab
    try:
        import reportlab  # noqa: F401
        _pdf_with_reportlab(path, lines, font_path)
        return path, "reportlab"
    except ImportError:
        pass
    except Exception as e:
        print(f"  [warn] reportlab 生成失败: {e}")

    # 2) fpdf2
    try:
        import fpdf  # noqa: F401
        _pdf_with_fpdf2(path, lines, font_path)
        return path, "fpdf2"
    except ImportError:
        pass
    except Exception as e:
        print(f"  [warn] fpdf2 生成失败: {e}")

    # 3) 纯标准库生成器（无第三方依赖，需要系统字体文件）
    if font_path:
        try:
            _pdf_with_stdlib(path, lines, font_path)
            return path, "stdlib(pure-python)"
        except Exception as e:
            print(f"  [warn] 纯标准库 PDF 生成失败: {e}")
    else:
        print("  [warn] 未找到可用的系统中文字体文件")

    # 4) 尝试自动安装 fpdf2
    print("  [info] 尝试自动安装 fpdf2 ...")
    try:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "fpdf2"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        _pdf_with_fpdf2(path, lines, font_path)
        return path, "fpdf2(auto-install)"
    except Exception as e:
        print(f"  [warn] 自动安装 fpdf2 失败: {e}")

    # 4) 回退：写成 .txt 文件
    txt_filename = os.path.splitext(filename)[0] + ".txt"
    txt_path = os.path.join(OUTPUT_DIR, txt_filename)
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"  [fallback] 已用 .txt 替代 PDF。如需真实 PDF，请运行: "
          f"{sys.executable} -m pip install fpdf2")
    return txt_path, "txt-fallback"


# --------------------------------------------------------------------------- #
# Word (.docx) 生成
# --------------------------------------------------------------------------- #
def write_docx(filename, lines):
    """使用 python-docx 生成 Word 文档。第一行作为标题，以「：」结尾的行加粗。"""
    from docx import Document

    path = os.path.join(OUTPUT_DIR, filename)
    doc = Document()

    title = lines[0].strip() if lines else filename
    doc.add_heading(title, level=0)

    for raw in lines[1:]:
        text = raw.rstrip("\n")
        if not text.strip():
            doc.add_paragraph("")
            continue
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        # 以「：」结尾的行视为小节标题，加粗显示
        if text.rstrip().endswith("："):
            run.bold = True

    doc.save(path)
    return path


# --------------------------------------------------------------------------- #
# Excel (.xlsx) 生成
# --------------------------------------------------------------------------- #
def write_xlsx(filename):
    """使用 openpyxl 生成 Excel：Sheet1 检验标准，Sheet2 判定规则。"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    path = os.path.join(OUTPUT_DIR, filename)
    wb = Workbook()

    # ---------- Sheet1: 检验标准 ----------
    ws = wb.active
    ws.title = "检验标准"

    headers = ["检验项目", "标准值", "公差", "检验方法"]
    data = [
        ["外径尺寸", "50.00mm", "±0.02mm", "千分尺测量"],
        ["内径尺寸", "30.00mm", "±0.01mm", "内径千分尺"],
        ["长度", "100.00mm", "±0.1mm", "游标卡尺"],
        ["表面粗糙度", "Ra1.6", "-", "粗糙度仪"],
        ["垂直度", "0.02mm", "-", "三坐标测量"],
    ]

    # 表头样式
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(name="微软雅黑", size=11, bold=True, color="FFFFFF")
    thin = Side(border_style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws.append(headers)
    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    for row in data:
        ws.append(row)

    # 数据单元格样式
    for row_idx in range(2, len(data) + 2):
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.border = border
            cell.alignment = Alignment(horizontal="center", vertical="center")

    # 列宽
    widths = [14, 12, 12, 16]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + i)].width = w

    # ---------- Sheet2: 判定规则 ----------
    ws2 = wb.create_sheet("判定规则")
    rules = [
        ["产品质量判定规则", ""],
        ["", ""],
        ["项目分类", "判定规则"],
        ["关键项目（外径、内径）", "任一不合格 → 产品不合格"],
        ["一般项目（长度、粗糙度）", "2 项以上不合格 → 产品不合格"],
        ["辅助项目（垂直度）", "仅记录，不判定"],
    ]
    for row in rules:
        ws2.append(row)

    # 标题行样式
    ws2["A1"].font = Font(name="微软雅黑", size=13, bold=True)
    ws2["A3"].font = Font(bold=True)
    ws2["B3"].font = Font(bold=True)
    ws2.column_dimensions["A"].width = 26
    ws2.column_dimensions["B"].width = 36

    wb.save(path)
    return path


# --------------------------------------------------------------------------- #
# Markdown (.md) 生成
# --------------------------------------------------------------------------- #
def write_md(filename):
    """直接写入 Markdown 内容。"""
    path = os.path.join(OUTPUT_DIR, filename)
    content = """# 设备操作经验记录

## 经验1：CNC机床换刀异常处理
- 现象：换刀时出现E04报警
- 原因：刀库位置传感器积灰
- 解决：清洁传感器后恢复正常
- 预防：每周清洁一次刀库传感器

## 经验2：液压系统压力波动
- 现象：压力表指针波动0.2-0.3MPa
- 原因：蓄能器氮气压力不足
- 解决：补充氮气至6MPa
- 预防：每季度检查蓄能器压力

## 经验3：加工表面粗糙度超标
- 现象：Ra值从1.6升至3.2
- 原因：刀具磨损或切削液浓度不足
- 解决：1.更换刀具 2.调整切削液浓度至8%
- 预防：每班次检查刀具磨损，每周检测切削液浓度
"""
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return path


# --------------------------------------------------------------------------- #
# 文档内容定义
# --------------------------------------------------------------------------- #
PDF_FAULT_MANUAL = [
    "设备故障代码手册",
    "",
    "E01 故障：电机过热",
    "- 原因：冷却系统故障或负载过大",
    "- 处理：1.停机检查冷却风扇 2.检查电机负载 3.清理散热器",
    "",
    "E02 故障：传感器异常",
    "- 原因：传感器接线松动或损坏",
    "- 处理：1.检查接线 2.更换传感器 3.重新校准",
    "",
    "E03 故障：气压不足",
    "- 原因：气路泄漏或空压机故障",
    "- 处理：1.检查气路密封 2.检查空压机 3.补充气压",
    "",
    "E04 故障：PLC通信中断",
    "- 原因：通信线缆断开或模块故障",
    "- 处理：1.检查通信线缆 2.重启通信模块 3.检查网络配置",
]

PDF_WAREHOUSE = [
    "仓库管理制度",
    "",
    "入库流程：",
    "1. 供应商送货 → 核对采购订单 → 质量检验 → 入库登记 → 上架",
    "2. 入库需在24小时内完成登记",
    "3. 不合格品进入退货区，48小时内处理",
    "",
    "出库流程：",
    "1. 领料申请 → 审批 → 拣货 → 复核 → 发料",
    "2. 生产用料需提前1天申请",
    "3. 紧急用料需车间主任签字",
    "",
    "盘点制度：",
    "1. 每日循环盘点（A类物料每日，B类每周，C类每月）",
    "2. 每季度全面盘点一次",
    "3. 盘点差异超过0.5%需查明原因",
]

DOCX_MAINTENANCE = [
    "设备保养计划规程",
    "",
    "日常保养（每日）：",
    "1. 开机前检查设备外观，清洁表面",
    "2. 检查油位、水位，不足时补充",
    "3. 检查气压是否正常（0.5-0.7MPa）",
    "4. 空载试运行3分钟，检查无异响",
    "",
    "周保养（每周）：",
    "1. 清洁导轨和丝杠，涂抹润滑脂",
    "2. 检查紧固件有无松动",
    "3. 检查电气连接是否可靠",
    "4. 校准传感器零点",
    "",
    "月保养（每月）：",
    "1. 更换液压油和润滑油",
    "2. 检查皮带张紧度，必要时更换",
    "3. 清洁电气柜，检查散热风扇",
    "4. 全面检查安全装置",
]

DOCX_CNC = [
    "数控加工工艺规程",
    "",
    "工序1：粗加工",
    "- 主轴转速：800 rpm",
    "- 进给速度：200 mm/min",
    "- 切削深度：3mm",
    "- 刀具：φ10立铣刀",
    "",
    "工序2：半精加工",
    "- 主轴转速：1200 rpm",
    "- 进给速度：150 mm/min",
    "- 切削深度：1mm",
    "- 刀具：φ8立铣刀",
    "",
    "工序3：精加工",
    "- 主轴转速：2000 rpm",
    "- 进给速度：100 mm/min",
    "- 切削深度：0.5mm",
    "- 刀具：φ6立铣刀",
    "",
    "注意事项：",
    "1. 粗加工后需冷却5分钟再进行半精加工",
    "2. 精加工前需校准刀具长度补偿",
]


# --------------------------------------------------------------------------- #
# 主流程
# --------------------------------------------------------------------------- #
def main():
    ensure_dir(OUTPUT_DIR)
    print(f"输出目录: {OUTPUT_DIR}")
    print("-" * 60)

    results = []

    # 1. 设备故障代码手册.pdf
    print("[1/6] 生成 设备故障代码手册.pdf ...")
    p, engine = write_pdf("设备故障代码手册.pdf", PDF_FAULT_MANUAL)
    results.append(("设备故障代码手册", p, engine))

    # 2. 设备保养计划规程.docx
    print("[2/6] 生成 设备保养计划规程.docx ...")
    p = write_docx("设备保养计划规程.docx", DOCX_MAINTENANCE)
    results.append(("设备保养计划规程", p, "python-docx"))

    # 3. 数控加工工艺规程.docx
    print("[3/6] 生成 数控加工工艺规程.docx ...")
    p = write_docx("数控加工工艺规程.docx", DOCX_CNC)
    results.append(("数控加工工艺规程", p, "python-docx"))

    # 4. 产品质量检验标准.xlsx
    print("[4/6] 生成 产品质量检验标准.xlsx ...")
    p = write_xlsx("产品质量检验标准.xlsx")
    results.append(("产品质量检验标准", p, "openpyxl"))

    # 5. 仓库管理制度.pdf
    print("[5/6] 生成 仓库管理制度.pdf ...")
    p, engine = write_pdf("仓库管理制度.pdf", PDF_WAREHOUSE)
    results.append(("仓库管理制度", p, engine))

    # 6. 设备操作经验记录.md
    print("[6/6] 生成 设备操作经验记录.md ...")
    p = write_md("设备操作经验记录.md")
    results.append(("设备操作经验记录", p, "markdown"))

    # 汇总
    print("-" * 60)
    print("生成完成！文件列表：")
    success = 0
    for name, path, engine in results:
        exists = os.path.exists(path)
        size = os.path.getsize(path) if exists else 0
        mark = "OK" if exists else "FAIL"
        if exists:
            success += 1
        rel = os.path.relpath(path, BACKEND_DIR)
        print(f"  [{mark}] {name:<14} ({engine:<18}) {size:>7} bytes  {rel}")
    print("-" * 60)
    print(f"成功: {success}/{len(results)}")

    if success < len(results):
        sys.exit(1)


if __name__ == "__main__":
    main()
